import asyncio
import threading
import time
import os
import platform
import requests
from DrissionPage import ChromiumPage, Chromium
from datetime import datetime
from docx import Document
from docx.shared import Inches

# 初始化会话，用于管理网络请求
session = requests.Session()

# 初始化当前评论元素
current_comment_element = None
# 存储视频链接的列表
video_url_list = []
# 用于存储文档内容的对象
document = Document()
# 总评论数量
total_comment_count = 0
# 地区关键词
region_keyword = None
# 关键词
keyword = None
# 是否插入图片
insert_image = True
# 缓存当前时间
current_time = datetime.now()
formatted_time = current_time.strftime("%Y-%m-%d %H-%M-%S")

async def download_image_async(url, name):
    """
    异步下载图片的函数
    """
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, download_image, url, name)

def download_image(url, name):
    """
    下载图片的函数
    :param url: 图片的URL
    :param name: 保存图片的文件名
    :return: 下载成功返回True，失败返回False
    """
    try:
        # 发送HTTP GET请求下载图片
        response = session.get(url)
        # 检查请求是否成功
        if response.status_code == 200:
            # 以二进制写入模式打开文件并写入图片内容
            with open(name, 'wb') as file:
                file.write(response.content)
            return True
        return False
    except Exception as e:
        return False

def get_clipboard_content():
    """
    获取剪贴板内容的函数，支持Windows、macOS和Linux系统
    :return: 剪贴板内容，如果获取失败返回空字符串
    """
    system_type = platform.system()  # 获取操作系统类型
    if system_type == "Windows":
        try:
            # 使用PowerShell命令获取剪贴板内容
            clipboard_content = os.popen('powershell -command "Get-Clipboard"').read().strip()
            return clipboard_content
        except Exception as e:
            print("无法获取剪贴板内容:", e)
            return ""
    elif system_type == "Darwin":  # macOS
        try:
            # 使用pbpaste命令获取剪贴板内容
            clipboard_content = os.popen('pbpaste').read().strip()
            return clipboard_content
        except Exception as e:
            print("无法获取剪贴板内容:", e)
            return ""
    elif system_type == "Linux":
        try:
            # 尝试使用xclip获取剪贴板内容
            clipboard_content = os.popen('xclip -selection clipboard -o').read().strip()
            return clipboard_content
        except Exception:
            try:
                # 尝试使用xsel获取剪贴板内容
                clipboard_content = os.popen('xsel --clipboard --output').read().strip()
                return clipboard_content
            except Exception as e:
                print("无法获取剪贴板内容:", e)
                return ""
    else:
        print("不支持的操作系统:", system_type)
        return ""

def get_comments(url):
    """
    获取指定视频链接的评论信息
    :param url: 视频链接
    """
    global document, total_comment_count, insert_image
    # 开始监听评论列表请求
    dp1.listen.start('comment/list/')
    # 访问视频链接
    dp1.get(url)
    get_comments_flag = True

    # 缓存视频标题
    try:
        video_title = dp1.ele('@@tag()=h1@@class=idrZUbq7').ele('@class=arnSiSbK').children()[0].children()[0].children()[0].children()[0].text
    except Exception:
        video_title = "未获取到标题"

    if insert_image:
        document.add_paragraph('')
        if os.path.exists('1.png'):
            document.add_picture('1.png', width=Inches(6))
        document.add_paragraph(f'视频链接:{url}')
        document.add_paragraph(f'视频标题:{video_title}')
        insert_image = False

    while get_comments_flag:
        try:
            # 等待评论列表请求响应
            response = dp1.listen.wait(timeout=5)
            if response:
                json_data = response.response.body
                if json_data is not None and 'comments' in json_data:
                    comments = json_data['comments']
                    if not comments:
                        get_comments_flag = False
                        break
                    for comment in comments:
                        if total_comment_count >= 500:
                            return
                        # 获取评论创建时间
                        create_time = comment.get('create_time', '')
                        date = str(datetime.fromtimestamp(create_time))
                        # 获取用户昵称
                        nickname = comment['user'].get('nickname', '')
                        # 获取抖音号
                        douyin_id = comment['user'].get('unique_id', '')
                        # 获取评论内容
                        comment_content = comment.get('text', '')
                        comment_element = None
                        nickname_element = None

                        # 定位评论元素
                        comment_box = dp1.ele('@class=Rwb9ssMc comment-mainContent')
                        for item in comment_box.children():
                            nickname_element = item.ele(f"@@tag()=span@@text()={nickname}", timeout=0.05)
                            if nickname_element:
                                comment_element = item
                                break

                        # 获取抖音号
                        if douyin_id == '':
                            if nickname_element:
                                nickname_element.click()
                                dp1.wait.load_start()
                                tab = dpBox.get_tabs()[0]
                                try:
                                    douyin_id = tab.ele('@class=OcCvtZ2a').text.split('抖音号：')[1]
                                except Exception:
                                    douyin_id = "未获取到抖音号"
                                document.add_paragraph(f'抖音号:{douyin_id}')
                                tab.close()
                                document.add_paragraph(f'抖昵称号:{nickname}')
                        else:
                            document.add_paragraph(f'抖音号:{douyin_id}')
                        document.add_paragraph(f'评论:{comment_content}')
                        document.add_paragraph(f'时间:{date}')
                        # 截图
                        if comment_element:
                            try:
                                comment_element.get_screenshot(name='0.png')
                                if os.path.exists('0.png'):
                                    document.add_picture('0.png', width=Inches(6))
                            except Exception as e:
                                print(f"截图失败: {e}")
                        document.add_paragraph('')

                        # 数据写入
                        data = {
                            '用户id': douyin_id,
                            '视频标题': video_title,
                            '昵称': nickname,
                            '地区': comment.get('ip_label', ''),
                            '日期': date,
                            '评论': comment_content,
                            '点赞数量': comment.get('digg_count', ''),
                            '视频链接': url
                        }
                        print(data)
                        total_comment_count += 1
                        print(f'评论数量{total_comment_count}')
                else:
                    get_comments_flag = False
            else:
                get_comments_flag = False
        except Exception as e:
            pass

        next_page = dp1.ele('css:.Rcc71LyU')
        if next_page:
            dp1.scroll.to_see(next_page)
        else:
            get_comments_flag = False

    # 处理完一个视频后保存文档
    if mode == '1':
        document.save(f'{region_keyword}{keyword}({formatted_time})-关键词评论.docx')
    elif mode == '2':
        document.save(f'{formatted_time}-关注评论.docx')

    # 重置插入图片标志，以便处理下一个视频
    insert_image = True

def get_video_links():
    """
    获取视频链接并处理评论信息
    """
    global keyword, region_keyword, document, insert_image

    if mode == '1':
        keyword = input("|请输入要查询的关键词->")
        # keyword = "房地产"
        region_keyword = input("|请输入要查询的地区关键词->")
        sort_option = input("|请选择排序方式(1.最新发布 2.综合排序)->")
        # 访问搜索页面
        dp.get(f'https://www.douyin.com/search/{region_keyword}{keyword}')

        # 点击视频标签
        click_video_flag = True
        video_button = dp.ele('@@class=gxIUdClv@@text()=视频', timeout=1)
        while click_video_flag and video_button:
            video_button.click()
            dp.wait.load_start()
            click_video_flag = False
            print('点击视频')
        if click_video_flag:
            print('视频点击失败')

        # 点击排序方式
        get_list_flag = True
        filter_button = dp.ele('@class=jjU9T0dQ')
        while get_list_flag and filter_button:
            try:
                dp.actions.move_to(filter_button)
                if sort_option == '1':
                    latest_button = dp.ele('@text()=最新发布')
                    if latest_button:
                        latest_button.click()
                        dp.wait.load_start()
                        get_list_flag = False
                        print('点击最新发布')
                    else:
                        print("最新发布按钮不在")
                elif sort_option == '2':
                    comprehensive_button = dp.ele('@text()=综合排序')
                    if comprehensive_button:
                        comprehensive_button.click()
                        dp.wait.load_start()
                        get_list_flag = False
                        print('点击综合排序')
                    else:
                        print("综合排序按钮不在")
                else:
                    print("无效的排序选项")
                    return
            except Exception as e:
                print('筛选按钮操作出错:', e)
        if get_list_flag:
            print('筛选按钮不在')

        print("开始获取进程")
        # 批量获取视频
        get_videos_flag = True
        # 缓存视频列表元素
        video_box = dp.ele('@@class=gZq36zrh PAjzsG5a@@tag()=ul')
        while get_videos_flag:
            for video_item in video_box.children():
                if total_comment_count >= 500:
                    return
                try:
                    # 缓存视频链接元素
                    video_link_ele = video_item.ele('@@class=hY8lWHgA _4furHfW@@tag()=a')
                    video_url = video_link_ele.link
                    if video_url not in video_url_list:
                        # 缓存视频封面元素
                        video_image_ele = video_item.ele('@class=VCzQd6LR zevBf7YE').ele('@tag()=img')
                        video_image_url = video_image_ele.link
                        threading.Thread(target=download_image, args=(video_image_url, '1.png')).start()
                        video_url_list.append(video_url)
                        get_comments(video_url)
                except Exception as e:
                    print(e)
            more = dp.ele('@text()=暂时没有更多了', timeout=0.5)
            if more:
                get_videos_flag = False
            else:
                dp.scroll.to_bottom()
            time.sleep(5)
    else:
        # 访问关注页面
        dp.get(f'https://www.douyin.com/follow')
        dp.wait.load_start()
        get_video_flag = True
        while get_video_flag:
            if total_comment_count >= 500:
                return
            try:
                share_button = dp.ele('@class=JPLz9DCE')
                # 判断是否为视频
                if share_button:
                    insert_image = True
                    # 点击主页
                    dp.ele('@class=B0JKdzQ8 KsoclCZj sVGJfEdt').click()
                    try:
                        video_image_ele = dp.ele('@@class=AI9Id9cO@@text()=播放中').parent().parent().ele('@@tag()=a@@class=uz1VJwFY TyuBARdT QLFBhWLd bHxHlnsh IdxE71f8').ele('@tag()=img')
                        video_image_url = video_image_ele.link
                        threading.Thread(target=download_image, args=(video_image_url, '1.png')).start()
                    except Exception as e:
                        print(f"获取视频图片链接失败: {e}")

                    dp.actions.move_to(share_button)
                    dp.ele('@text()=复制链接').click()
                    copied_text = get_clipboard_content()
                    try:
                        copied_text = copied_text.split('/ ')
                        copied_text = copied_text[len(copied_text) - 2]
                        copied_text = copied_text.split(' ')
                        video_url = copied_text[len(copied_text) - 1]
                        get_comments(video_url)
                    except IndexError:
                        print("解析复制链接失败")

                # 点击下一个视频
                button_box = dp.ele('@class=xgplayer-playswitch-tab')
                next_button = button_box.ele('@class=xgplayer-playswitch-next')
                next_button.click()
                time.sleep(5)
            except Exception as e:
                dp.refresh()


# 确认登陆信息
dpBox = Chromium(6666)
dp = dpBox.new_tab()
dp1 = dpBox.new_tab()
dp.get('https://v.douyin.com/')
input("|确认登陆信息(按回车继续)")
mode = input("|请输入要进行的模式(1.关键词 2.关注)->")

try:
    get_video_links()
except KeyboardInterrupt:
    print("\n检测到 Ctrl+C，程序正在退出...")
    if mode == '1':
        document.save(f'{region_keyword}{keyword}({formatted_time})-关键词评论.docx')
    elif mode == '2':
        document.save(f'{formatted_time}-关注评论.docx')