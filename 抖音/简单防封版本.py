import asyncio
import threading
import time
import os
import platform
import requests
import random
import signal
from DrissionPage import ChromiumPage, Chromium
from datetime import datetime
from docx import Document, document
from docx.shared import Inches

# 全局停止标志
stop_flag = False

# 初始化会话，用于管理网络请求
session = requests.Session()

# 存储视频链接的列表
video_url_list = []
# 总评论数量
total_comment_count = 0
# 每个视频的评论数量
comments_per_video = 30
# 地区关键词
region_keyword = None
# 关键词
keyword = None
# 是否插入图片
insert_image = True
# 当前时间
current_time = datetime.now()
formatted_time = current_time.strftime("%Y-%m-%d %H-%M-%S")

# 定义请求延迟范围（秒）
MIN_DELAY = 0.5
MAX_DELAY = 2
# 最大重试次数
MAX_RETRIES = 3

# 随机请求头列表
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
    "Mozilla/5.0 (iPhone; CPU iPhone OS 14_6 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.1.1 Mobile/15E148 Safari/604.1",
    # 添加更多不同的 User-Agent
]


def signal_handler(sig, frame):
    global stop_flag
    print("\n接收到停止信号，正在保存数据...")
    stop_flag = True


signal.signal(signal.SIGINT, signal_handler)


def download_image(url, name):
    """
    下载图片的函数
    :param url: 图片的URL
    :param name: 保存图片的文件名
    :return: 下载成功返回True，失败返回False
    """
    retries = 0
    while retries < MAX_RETRIES and not stop_flag:
        try:
            headers = {'User-Agent': random.choice(USER_AGENTS)}
            response = session.get(url, headers=headers)
            if response.status_code == 200:
                with open(name, 'wb') as file:
                    file.write(response.content)
                return True
            else:
                print(f"图片下载失败，状态码: {response.status_code}，重试 {retries + 1}/{MAX_RETRIES}")
        except Exception as e:
            print(f"图片下载出错: {e}，重试 {retries + 1}/{MAX_RETRIES}")
        retries += 1
        time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))
    return False


def _is_valid_png(file_path):
    """验证PNG文件有效性"""
    try:
        with open(file_path, 'rb') as f:
            header = f.read(8)
            return header == b'\x89PNG\r\n\x1a\n'
    except Exception:
        return False


def get_clipboard_content():
    """
    获取剪贴板内容的函数，支持Windows、macOS和Linux系统
    :return: 剪贴板内容，如果获取失败返回空字符串
    """
    system_type = platform.system()  # 获取操作系统类型
    if system_type == "Windows":
        try:
            # 使用PowerShell命令获取剪贴板内容
            clipboard_content = os.popen('powershell -command "Get-Clipboard"').read().strip()
            return clipboard_content
        except Exception as e:
            print("无法获取剪贴板内容:", e)
            return ""
    elif system_type == "Darwin":  # macOS
        try:
            # 使用pbpaste命令获取剪贴板内容
            clipboard_content = os.popen('pbpaste').read().strip()
            return clipboard_content
        except Exception as e:
            print("无法获取剪贴板内容:", e)
            return ""
    elif system_type == "Linux":
        try:
            # 尝试使用xclip获取剪贴板内容
            clipboard_content = os.popen('xclip -selection clipboard -o').read().strip()
            return clipboard_content
        except Exception:
            try:
                # 尝试使用xsel获取剪贴板内容
                clipboard_content = os.popen('xsel --clipboard --output').read().strip()
                return clipboard_content
            except Exception as e:
                print("无法获取剪贴板内容:", e)
                return ""
    else:
        print("不支持的操作系统:", system_type)
        return ""


def get_comments(url, document):
    """
    获取指定视频链接的评论信息
    :param url: 视频链接
    :param document: 文档对象
    """
    global total_comment_count, comments_per_video, insert_image, stop_flag
    # 开始监听评论列表请求
    dp1 = dpBox.new_tab()
    dp1.listen.start('comment/list/')
    retries = 0
    while retries < MAX_RETRIES and not stop_flag:
        try:
            time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
            headers = {'User-Agent': random.choice(USER_AGENTS)}
            dp1.headers = headers
            # 访问视频链接
            dp1.get(url)
            # 模拟随机滚动
            scroll_distance = random.randint(100, 500)
            dp1.scroll.to(0, scroll_distance)
            time.sleep(random.uniform(0.2, 0.5))  # 滚动后稍作停留
            break
        except Exception as e:
            print(f"访问视频链接出错: {e}，重试 {retries + 1}/{MAX_RETRIES}")
            retries += 1
            time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))
    if retries == MAX_RETRIES or stop_flag:
        print("访问视频链接失败，放弃该视频")
        dp1.quit()
        return

    current_comment_count = 0
    get_comments_flag = True

    while get_comments_flag and not stop_flag:
        try:
            # 等待评论列表请求响应
            response = dp1.listen.wait(timeout=5)
            if response:
                json_data = response.response.body
                if json_data is not None and 'comments' in json_data:
                    comments = json_data['comments']
                    if comments is not None:
                        for comment in comments:
                            if current_comment_count >= comments_per_video or stop_flag:
                                get_comments_flag = False
                                break
                            # 获取评论创建时间
                            create_time = comment.get('create_time', '')
                            date = str(datetime.fromtimestamp(create_time))
                            # 获取用户昵称
                            nickname = comment['user'].get('nickname', '')
                            # 获取抖音号
                            douyin_id = comment['user'].get('unique_id', '')
                            # 获取评论内容
                            comment_content = comment.get('text', '')
                            comment_element = None
                            nickname_element = None

                            # 定位评论元素
                            comment_box = dp1.ele('@class=Rwb9ssMc comment-mainContent')
                            for item in comment_box.children():
                                # 优化后的昵称定位
                                nickname_elements = item.eles('@class=DyFp1rqB', timeout=0.05)
                                if nickname_elements:
                                    nickname_element = nickname_elements[0]
                                    comment_element = item
                                    break

                            if insert_image:
                                if os.path.exists('1.png') and _is_valid_png('1.png'):
                                    document.add_paragraph('')
                                    document.add_picture('1.png', width=Inches(6))
                                # 优化后的标题获取
                                try:
                                    video_title = dp1.ele('@class=h1', timeout=5).text
                                except Exception:
                                    video_title = "未获取到标题"
                                document.add_paragraph(f'视频链接:{url}')
                                document.add_paragraph(f'视频标题:{video_title}')
                                insert_image = False

                            # 获取抖音号
                            if douyin_id == '':
                                if nickname_element:
                                    time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                                    nickname_element.click()
                                    dp1.wait.load_start()
                                    tab = dpBox.get_tabs()[0]
                                    try:
                                        douyin_id = tab.ele('@class=OcCvtZ2a').text.split('抖音号：')[1]
                                    except Exception:
                                        douyin_id = "未获取到抖音号"
                                    document.add_paragraph(f'抖音号:{douyin_id}')
                                    tab.close()
                                    document.add_paragraph(f'抖昵称号:{nickname}')
                            else:
                                document.add_paragraph(f'抖音号:{douyin_id}')
                            document.add_paragraph(f'评论:{comment_content}')
                            document.add_paragraph(f'时间:{date}')
                            # 优化后的截图
                            if comment_element:
                                try:
                                    comment_element.get_screenshot(name='0.png', captureBeyondViewport=True)
                                    if os.path.exists('0.png') and _is_valid_png('0.png'):
                                        document.add_picture('0.png', width=Inches(6))
                                except Exception as e:
                                    print(f"截图失败: {e}")
                                    dp1.get_screenshot(name='0.png')
                                    if os.path.exists('0.png') and _is_valid_png('0.png'):
                                        document.add_picture('0.png', width=Inches(6))
                            document.add_paragraph('')

                            # 数据写入
                            data = {
                                '用户id': douyin_id,
                                '视频标题': video_title,
                                '昵称': nickname,
                                '地区': comment.get('ip_label', ''),
                                '日期': date,
                                '评论': comment_content,
                                '点赞数量': comment.get('digg_count', ''),
                                '视频链接': url
                            }
                            print(data)
                            total_comment_count += 1
                            print(f'评论数量{total_comment_count}')

                            current_comment_count += 1
                            time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                            document.save(f'{region_keyword}{keyword}({formatted_time})-关键词评论.docx')
                else:
                    current_comment_count += comments_per_video
            else:
                current_comment_count += comments_per_video
        except Exception as e:
            print(f"处理评论时出现异常: {e}")
            import traceback
            traceback.print_exc()
        if current_comment_count >= comments_per_video or stop_flag:
            get_comments_flag = False
            break
        else:
            next_page = dp1.ele('css:.Rcc71LyU')
            if next_page:
                dp1.scroll.to_see(next_page)
    dp1.quit()


def process_video(url, document):
    """
    处理单个视频的函数，用于多线程调用
    """
    global stop_flag
    while not stop_flag:
        video_image_ele = dp.ele('@class=VCzQd6LR zevBf7YE').ele('@tag()=img')
        video_image_url = video_image_ele.link
        if download_image(video_image_url, '1.png'):
            get_comments(url, document)
            break
        time.sleep(1)


def get_video_links():
    """
    获取视频链接并处理评论信息
    """
    global keyword, region_keyword, comments_per_video, insert_image, stop_flag
    document = Document()

    if mode == '1':
        keyword = input("|请输入要查询的关键词->")
        region_keyword = input("|请输入要查询的地区关键词->")
        comments_per_video = int(input("|请输入每个视频爬取的最小评论数量(建议30)->"))
        # 访问搜索页面
        retries = 0
        while retries < MAX_RETRIES and not stop_flag:
            try:
                time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                headers = {'User-Agent': random.choice(USER_AGENTS)}
                dp.headers = headers
                dp.get(f'https://www.douyin.com/search/{region_keyword}{keyword}')
                # 模拟随机滚动
                scroll_distance = random.randint(100, 500)
                dp.scroll.to_location(0, scroll_distance)
                time.sleep(random.uniform(0.2, 0.5))  # 滚动后稍作停留
                break
            except Exception as e:
                print(f"访问搜索页面出错: {e}，重试 {retries + 1}/{MAX_RETRIES}")
                retries += 1
                time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))
        if retries == MAX_RETRIES or stop_flag:
            print("访问搜索页面失败，退出程序")
            return

        # 点击视频标签
        click_video_flag = True
        while click_video_flag and not stop_flag:
            video_button = dp.ele('@@class=gxIUdClv@@text()=视频', timeout=1)
            if video_button:
                time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                video_button.click()
                dp.wait.load_start()
                click_video_flag = False
                print('点击视频')
            else:
                dp.refresh()
                print('视频点击失败')

        # 点击最新发布
        get_list_flag = True
        while get_list_flag and not stop_flag:
            try:
                get_list_flag = False
                filter_button = dp.ele('@class=jjU9T0dQ')
                if filter_button:
                    dp.actions.move_to(filter_button)
                    latest_button = dp.ele('@text()=最新发布')
                    if latest_button:
                        time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                        latest_button.click()
                        dp.wait.load_start()
                        print('点击最新发布')
                    else:
                        get_list_flag = True
                        dp.refresh()
                        print("最新发布按钮不在")
                else:
                    dp.refresh()
                    get_list_flag = True
                    print('筛选按钮不在')
            except Exception as e:
                dp.refresh()
                get_list_flag = True
                print('筛选按钮不在')

        print("开始获取进程")
        # 批量获取视频
        get_videos_flag = True
        # 缓存视频列表元素
        video_box = dp.ele('@@class=gZq36zrh PAjzsG5a@@tag()=ul')
        threads = []
        while get_videos_flag and not stop_flag:
            for video_item in video_box.children():
                # 缓存视频链接元素
                video_link_ele = video_item.ele('@@class=hY8lWHgA _4furHfW@@tag()=a')
                video_url = video_link_ele.link
                if video_url not in video_url_list:
                    video_url_list.append(video_url)
                    time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                    thread = threading.Thread(target=process_video, args=(video_url, document))
                    threads.append(thread)
                    thread.start()
            more = dp.ele('@text()=暂时没有更多了', timeout=0.5)
            if more:
                get_videos_flag = False
            else:
                dp.scroll.to_bottom()
            time.sleep(5)

        # 等待所有线程完成
        for thread in threads:
            if thread.is_alive():
                thread.join()

        if not stop_flag:
            document.save(f'{region_keyword}{keyword}({formatted_time})-关键词评论.docx')
    else:
        comments_per_video = int(input("|请输入每个视频爬取的最小评论数量(建议30)->"))
        # 访问关注页面
        retries = 0
        while retries < MAX_RETRIES and not stop_flag:
            try:
                time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                headers = {'User-Agent': random.choice(USER_AGENTS)}
                dp.headers = headers
                dp.get(f'https://www.douyin.com/follow')
                # 模拟随机滚动
                scroll_distance = random.randint(100, 500)
                dp.scroll.to(0, scroll_distance)
                time.sleep(random.uniform(0.2, 0.5))  # 滚动后稍作停留
                break
            except Exception as e:
                print(f"访问关注页面出错: {e}，重试 {retries + 1}/{MAX_RETRIES}")
                retries += 1
                time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))
        if retries == MAX_RETRIES or stop_flag:
            print("访问关注页面失败，退出程序")
            return

        dp.wait.load_start()
        get_video_flag = True
        threads = []
        while get_video_flag and not stop_flag:
            try:
                share_button = dp.ele('@class=JPLz9DCE')
                # 判断是否为视频
                if share_button:
                    insert_image = True
                    # 点击主页
                    time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                    dp.ele('@class=B0JKdzQ8 KsoclCZj sVGJfEdt').click()
                    try:
                        video_image_ele = dp.ele('@@class=AI9Id9cO@@text()=播放中').parent().parent().ele('@@tag()=a@@class=uz1VJwFY TyuBARdT QLFBhWLd bHxHlnsh IdxE71f8').ele('@tag()=img')
                        video_image = video_image_ele.link
                    except Exception as e:
                        video_image = ''
                    if video_image:
                        dp.actions.move_to(share_button)
                        time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                        dp.ele('@text()=复制链接').click()
                        copied_text = get_clipboard_content()
                        copied_text = copied_text.split('/ ')
                        copied_text = copied_text[len(copied_text) - 2]
                        copied_text = copied_text.split(' ')
                        video_url = copied_text[len(copied_text) - 1]
                        time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                        thread = threading.Thread(target=process_video, args=(video_url, document))
                        threads.append(thread)
                        thread.start()

                # 点击下一个视频
                button_box = dp.ele('@class=xgplayer-playswitch-tab')
                next_button = button_box.ele('@class=xgplayer-playswitch-next')
                time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))  # 随机延迟
                next_button.click()
                time.sleep(5)
            except Exception as e:
                dp.refresh()

        # 等待所有线程完成
        for thread in threads:
            if thread.is_alive():
                thread.join()

        if not stop_flag:
            document.save(f'{formatted_time}-关注评论.docx')


# 确认登陆信息
dpBox = Chromium(6666)
dp = dpBox.new_tab()
dp1 = dpBox.new_tab()
dp.get('https://v.douyin.com/')
input("|确认登陆信息(按回车继续)")
mode = input("|请输入要进行的模式(1.关键词 2.关注)->")

try:
    get_video_links()
except Exception as e:
    print(f"程序异常终止: {e}")
finally:
    print("正在清理资源...")
    dpBox.close_tabs()
    if 'document' in locals():
        try:
            if mode == '1':
                document.save(f'{region_keyword}{keyword}({formatted_time})-关键词评论.docx')
            elif mode == '2':
                document.save(f'{formatted_time}-关注评论.docx')
        except Exception as e:
            print(f"保存文档失败: {e}")