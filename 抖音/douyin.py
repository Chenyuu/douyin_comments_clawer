from DrissionPage import ChromiumPage,Chromium
from datetime import datetime
import time
import os
from docx import Document
from docx.shared import Inches
import platform
import requests
# 当前评论的元素
plBox=None
# 视频链接
videoUrlList=[]
# 存储
doc = Document()
# 总评论数量
allPlNum=0
# 每个视频的评论数量
plNum=30
# 地区关键词
dqgjc=None
# 关键词
gjc=None
# 是否放入图片
setImg=True
# 当前时间
current_time = datetime.now()
nowTime=current_time.strftime("%Y-%m-%d %H-%M-%S")

def downImg(url,name):
    try:
        # 发送HTTP GET请求，下载图片
        response = requests.get(url)

        # 检查请求是否成功
        if response.status_code == 200:
            # 图片的文件名（可以从URL中提取，或者自定义）
            filename = name
            
            # 以二进制写入模式打开文件，并将下载的内容写入文件
            with open(filename, 'wb') as f:
                f.write(response.content)
            return True
        else:
            return False
    except Exception as e:
        return False

def get_clipboard_content():
    system = platform.system()  # 获取操作系统类型
    if system == "Windows":
        try:
            clipboard_content = os.popen('powershell -command "Get-Clipboard"').read().strip()
            return clipboard_content
        except Exception as e:
            print("无法获取剪贴板内容:", e)
            return ""
    elif system == "Darwin":  # macOS
        try:
            clipboard_content = os.popen('pbpaste').read().strip()
            return clipboard_content
        except Exception as e:
            print("无法获取剪贴板内容:", e)
            return ""
    elif system == "Linux":
        try:
            # 尝试使用 xclip
            clipboard_content = os.popen('xclip -selection clipboard -o').read().strip()
            return clipboard_content
        except Exception:
            try:
                # 尝试使用 xsel
                clipboard_content = os.popen('xsel --clipboard --output').read().strip()
                return clipboard_content
            except Exception as e:
                print("无法获取剪贴板内容:", e)
                return ""
    else:
        print("不支持的操作系统:", system)
        return ""


def getComment(url):
    global doc
    global allPlNum
    global plNum
    global setImg

    dp1.listen.start('comment/list/')
    dp1.get(url)
    nowPlNum=0
    getPlRunBtn=True

    while getPlRunBtn:
        try:
            resp=dp1.listen.wait(timeout=5)
            if(resp!=False):
                json_data=resp.response.body
                if json_data is not None:
                    # 检查'comments'键是否存在于json_data字典中，并且它的值不是空列表
                    if 'comments' in json_data and json_data['comments']:
                        comments = json_data['comments']
                        for index in comments:
                            create_time=index.get('create_time','')
                            date=str(datetime.fromtimestamp(create_time))
                            nickname=index['user'].get('nickname', '')
                            dyh=index['user'].get('unique_id','')
                            contentpl=index.get('text','')
                            plele=None
                            nickNameEle=None

                            # 定位名字
                            plBox=dp1.ele('@class=Rwb9ssMc comment-mainContent')
                            for item in plBox.children():
                                nickNameEle=item.ele(f"@@tag()=span@@text()={nickname}",timeout=0.05)
                                if nickNameEle:
                                        plele=item
                                        break

                            if setImg:
                                doc.add_paragraph('')
                                doc.add_picture('1.png', width=Inches(6))
                                # 获取视频标题
                                videoContent=dp1.ele('@@tag()=h1@@class=idrZUbq7').ele('@class=arnSiSbK').children()[0].children()[0].children()[0].children()[0].text
                                doc.add_paragraph(f'视频链接:{url}')
                                doc.add_paragraph(f'视频标题:{videoContent}')
                                setImg=False

                            # 获取抖音号
                            if dyh=='':
                                nickNameEle.click()
                                dp1.wait.load_start()
                                ptab=dpBox.get_tabs()[0]
                                dyh=ptab.ele('@class=OcCvtZ2a').text.split('抖音号：')[1]
                                doc.add_paragraph(f'抖音号:{dyh}')
                                ptab.close()
                                doc.add_paragraph(f'抖昵称号:{nickname}')
                            else:
                                doc.add_paragraph(f'抖音号:{dyh}')
                            doc.add_paragraph(f'评论:{contentpl}')
                            doc.add_paragraph(f'时间:{date}')
                            # 截图
                            plele.get_screenshot(name='0.png')
                            doc.add_picture('0.png', width=Inches(6))
                            doc.add_paragraph('')
                            if mode=='1':
                                doc.save(f'{dqgjc}{gjc}({nowTime})-关键词评论.docx')
                            elif mode=='2':
                                doc.save(f'{nowTime}-关注评论.docx')

                            # 数据写入
                            dit={
                                '用户id':dyh,
                                '视频标题':videoContent,
                                '昵称':nickname,
                                '地区':index.get('ip_label',''),
                                '日期':date,
                                '评论':contentpl,
                                '点赞数量':index.get('digg_count',''),
                                '视频链接':url
                            }
                            # csv_writer.writerow(dit)
                            print(dit)
                            allPlNum+=1
                            print(f'评论数量{allPlNum}')

                            nowPlNum+=1
                    else:
                        pass
                        nowPlNum+=plNum
                else:
                    pass
                    nowPlNum+=plNum
            else:
                pass
                nowPlNum+=plNum
        except Exception as e:
            # print(f"发生错误: {e}")
            # input()
            pass
        if nowPlNum>=plNum:
            getPlRunBtn=False
            break
        else:
            next_page=dp1.ele('css:.Rcc71LyU')
            if(next_page):
                dp1.scroll.to_see(next_page)


def getLink():
    global gjc
    global dqgjc
    global plNum
    global doc
    global setImg

    if mode=='1':
        gjc=input("|请输入要查询的关键词->")
        dqgjc=input("|请输入要查询的地区关键词->")
        plNum=int(input("|请输入每个视频爬取的最小评论数量(建议30)->"))
        # gjc='房地产'
        # dqgjc='广东'
        # plNum=30
        dp.get(f'https://www.douyin.com/search/{dqgjc}{gjc}')

        # 点击视频
        clickVideoRunBtn=True
        while clickVideoRunBtn:
            videoBtn=dp.ele('@@class=gxIUdClv@@text()=视频',timeout=1)
            if videoBtn:
                videoBtn.click()
                dp.wait.load_start()
                clickVideoRunBtn=False
                print('点击视频')
            else:
                dp.refresh()
                print('视频点击失败')

        # 点击最新发布
        getListBtn=True
        while getListBtn:
            try:
                getListBtn=False
                shaixuanBtn=dp.ele('@class=jjU9T0dQ')
                if shaixuanBtn:
                    dp.actions.move_to(shaixuanBtn)
                    zxfbBtn=dp.ele('@text()=最新发布')
                    if zxfbBtn:
                        zxfbBtn.click()
                        dp.wait.load_start()
                        print('点击最新发布')
                    else:
                        getListBtn=True
                        dp.refresh()
                        print("最新发布按钮不在")

                else:
                    dp.refresh()
                    getListBtn=True
                    print('筛选按钮不在')
            except Exception as e:
                dp.refresh()
                getListBtn=True
                print('筛选按钮不在')

        print("开始获取进程")
        # 批量获取视频
        getVideosRunBtn=True
        while getVideosRunBtn:
            videoBox=dp.ele('@@class=gZq36zrh PAjzsG5a@@tag()=ul')
            for videoItem in videoBox.children():
                getVideoRunBtn=True
                while getVideoRunBtn:
                    try:
                        videoUrl=videoItem.ele('@@class=hY8lWHgA _4furHfW@@tag()=a').link
                        if not(videoUrl in videoUrlList):
                            # 封面
                            videoImgUrl=videoItem.ele('@class=VCzQd6LR zevBf7YE').ele('@tag()=img').link
                            downImg(videoImgUrl,'1.png')
                            setImg=True
                            videoUrlList.append(videoUrl)
                            getComment(videoUrl)
                        getVideoRunBtn=False
                    except Exception as e:
                        print(e)
                        getVideoRunBtn=True
            # videos=dp.eles('@@class=hY8lWHgA _4furHfW@@tag()=a')
            # for videoItem in videos:
            #     if not(videoItem.link in videoUrlList):
            #         videoUrlList.append(videoItem.link)
            #         getComment(videoItem.link)
            more=dp.ele('@text()=暂时没有更多了',timeout=0.5)
            if more:
                getVideosRunBtn=False
            else:
                dp.scroll.to_bottom()
            time.sleep(5)
    else:
        plNum=int(input("|请输入每个视频爬取的最小评论数量(建议30)->"))
        # plNum=30
        dp.get(f'https://www.douyin.com/follow')
        dp.wait.load_start()
        getVideoBtn=True
        while getVideoBtn:
            try:
                fxBtn=dp.ele('@class=JPLz9DCE')
                # 判断是不是视频
                if fxBtn:
                    setImg=True
                    # 点击主页
                    dp.ele('@class=B0JKdzQ8 KsoclCZj sVGJfEdt').click()
                    try:
                        videoImg=dp.ele('@@class=AI9Id9cO@@text()=播放中').parent().parent().ele('@@tag()=a@@class=uz1VJwFY TyuBARdT QLFBhWLd bHxHlnsh IdxE71f8').ele('@tag()=img').link
                    except Exception as e:
                        videoImg=''
                    if videoImg:
                        downImg(videoImg,'1.png')
                        
                    dp.actions.move_to(fxBtn)
                    dp.ele('@text()=复制链接').click()
                    copied_text = get_clipboard_content()
                    copied_text=copied_text.split('/ ')
                    copied_text=copied_text[len(copied_text)-2]
                    copied_text=copied_text.split(' ')
                    videoUrl=copied_text[len(copied_text)-1]
                    getComment(videoUrl)
                
                # 点击下一个视频
                btnBox=dp.ele('@class=xgplayer-playswitch-tab')
                nextBtn=btnBox.ele('@class=xgplayer-playswitch-next')
                nextBtn.click()
                time.sleep(5)
            except Exception as e:
                dp.refresh()
        
# 确认登陆信息
dpBox=Chromium(6666)
dp=dpBox.new_tab()
dp1=dpBox.new_tab()
dp.get('https://v.douyin.com/')
input("|确认登陆信息(按回车继续)")
mode=input("|请输入要进行的模式(1.关键词 2.关注)->")

try:
    getLink()
except KeyboardInterrupt:
    print("\n检测到 Ctrl+C，程序正在退出...")
    doc.save(f'{dqgjc}{gjc}({nowTime})/评论.docx')
